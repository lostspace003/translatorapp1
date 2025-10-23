from pathlib import Path
from typing import Literal
from io import BytesIO

from docx import Document
from PyPDF2 import PdfReader
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem, PageBreak, Table, TableStyle, LongTable
from reportlab.lib.units import cm
from reportlab.lib import colors

import re

_BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")

def strip_markdown(text: str) -> str:
    """Remove simple markdown (currently **bold** only) for TXT export."""
    return _BOLD_RE.sub(lambda m: m.group(1), text)

def _split_bold_segments(text: str):
    """Split a string into [(segment, is_bold), ...] by **bold** markers."""
    parts = []
    last = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > last:
            parts.append((text[last:m.start()], False))
        parts.append((m.group(1), True))
        last = m.end()
    if last < len(text):
        parts.append((text[last:], False))
    return parts

FileKind = Literal["pdf", "docx", "txt", "csv", "xlsx", "image", "unknown"]

def detect_type(ext: str) -> FileKind:
    ext = ext.lower()
    if ext == ".pdf":
        return "pdf"
    if ext in {".docx"}:  # legacy .doc not supported; please convert to .docx
        return "docx"
    if ext in {".txt"}:
        return "txt"
    if ext in {".csv"}:
        return "csv"
    if ext in {".xlsx"}:  # legacy .xls not supported; please convert to .xlsx
        return "txt"
    if ext in {".png", ".jpg", ".jpeg"}:
        return "image"
    return "unknown"

def extract_text_from_pdf(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))
    texts = []
    for page in reader.pages:
        try:
            texts.append(page.extract_text() or "")
        except Exception:
            pass
    return "\n\n".join(t.strip() for t in texts if t and t.strip())

def extract_text_from_docx(data: bytes) -> str:
    f = BytesIO(data)
    doc = Document(f)
    return "\n\n".join(p.text for p in doc.paragraphs)

def extract_text_from_txt(data: bytes) -> str:
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("latin-1", errors="ignore")

def save_txt(text: str, path: Path) -> None:
    # strip markdown markers for plain text files
    path.write_text(strip_markdown(text), encoding="utf-8")


def save_docx(text: str, path: Path) -> None:
    doc = Document()
    paragraphs = text.split("\n\n")
    for para in paragraphs:
        lines = para.split("\n")
        for li, line in enumerate(lines):
            p = doc.add_paragraph()
            for seg, is_bold in _split_bold_segments(line):
                if not seg:
                    continue
                r = p.add_run(seg)
                if is_bold:
                    r.bold = True
        # blank line between paragraphs
    doc.save(path)



def extract_text_from_csv(data: bytes) -> str:
    """Decode CSV bytes and render as pipe-delimited text (one row per line)."""
    import csv
    from io import StringIO
    # try UTF-8, fallback latin-1
    try:
        txt = data.decode("utf-8")
    except UnicodeDecodeError:
        txt = data.decode("latin-1", errors="ignore")
    reader = csv.reader(StringIO(txt))
    lines = []
    for row in reader:
        safe = ["" if v is None else str(v) for v in row]
        lines.append(" | ".join(safe).rstrip())
    return "\n".join(lines)

def extract_text_from_xlsx(data: bytes) -> str:
    """Read .xlsx and flatten each worksheet as pipe-delimited rows with sheet headers."""
    from openpyxl import load_workbook
    wb = load_workbook(filename=BytesIO(data), data_only=True, read_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"===== Sheet: {ws.title} =====")
        for row in ws.iter_rows(values_only=True):
            cells = ["" if v is None else str(v) for v in row]
            parts.append(" | ".join(cells).rstrip())
        parts.append("")
    return "\n".join(parts).strip()




def save_csv_from_flat(text: str, path: Path) -> None:
    """Save pipe-delimited lines to CSV.
    - Splits each non-empty line on ' | '
    - Ignores sheet headers like '===== Sheet: Name =====' (content-only)
    """
    import csv, re
    rows = []
    for line in text.split("\n"):
        s = line.strip()
        if not s:
            continue
        if re.match(r"^=+\s*Sheet:\s*.+?\s*=+$", s, flags=re.I):
            # sheet header; skip in CSV
            continue
        cells = [c.strip() for c in line.split(" | ")]
        rows.append(cells)
    with path.open("w", encoding="utf-8", newline="") as f:
        writer = csv.writer(f)
        writer.writerows(rows)

def save_xlsx_from_flat(text: str, path: Path) -> None:
    """Save flattened text to XLSX.
    - Creates new sheet when encountering '===== Sheet: Name ====='
    - Otherwise splits lines on ' | ' into cells
    """
    from openpyxl import Workbook
    import re
    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    first_sheet_used = False

    def new_sheet(name: str):
        nonlocal ws, first_sheet_used
        if not first_sheet_used:
            # Rename first sheet
            ws.title = name or "Sheet1"
            first_sheet_used = True
        else:
            ws = wb.create_sheet(title=name or "Sheet")

    for line in text.split("\n"):
        s = line.strip()
        if not s:
            continue
        m = re.match(r"^=+\s*Sheet:\s*(.+?)\s*=+$", s, flags=re.I)
        if m:
            new_sheet(m.group(1))
            continue
        cells = [c.strip() for c in line.split(" | ")]
        ws.append(cells)
    wb.save(path)



def save_pdf(text: str, path: Path) -> None:
    """Generate a nicely formatted PDF using ReportLab Platypus.
    - Converts **bold**/*italic*
    - Headings: #, ##, ###
    - Lists: -, *, 1.
    - Quotes: > prefix
    - Tables: lines with pipes (|) and optional '===== Sheet: Name =====' headers
    - Page breaks: a block exactly '---'
    """
    from xml.sax.saxutils import escape as xml_escape
    import re

    styles = getSampleStyleSheet()
    base = styles["BodyText"]
    base.leading = 14
    h1 = styles["Heading1"]
    h2 = styles["Heading2"]
    h3 = styles["Heading3"]
    quote = ParagraphStyle("Quote", parent=base, leftIndent=1*cm, italic=True, textColor="#333333")

    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm
    )

    def inline_markup(s: str) -> str:
        s = xml_escape(s)
        s = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", s)
        s = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"<i>\1</i>", s)
        s = s.replace("\n", "<br/>")
        return s

    def is_list_block(lines):
        stripped_lines = [ln for ln in lines if ln.strip()]
        if not stripped_lines:
            return False
        ok = True
        for ln in stripped_lines:
            s = ln.strip()
            if not (s.startswith("-") or s.startswith("*") or re.match(r"\d+\.", s)):
                ok = False
                break
        return ok

    def is_table_block(lines):
        candidates = [ln for ln in lines if '|' in ln]
        return len(candidates) >= max(2, int(0.6 * len([ln for ln in lines if ln.strip()])))

    def table_from_lines(lines):
        rows = []
        for ln in lines:
            if '|' not in ln:
                continue
            parts = [c.strip() for c in ln.split('|')]
            rows.append(parts)
        if not rows:
            return None
        tbl = Table(rows, hAlign='LEFT')
        style = TableStyle([
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('BACKGROUND', (0,0), (-1,0), colors.whitesmoke),
            ('GRID', (0,0), (-1,-1), 0.25, colors.grey),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('LEFTPADDING', (0,0), (-1,-1), 6),
            ('RIGHTPADDING', (0,0), (-1,-1), 6),
            ('TOPPADDING', (0,0), (-1,-1), 4),
            ('BOTTOMPADDING', (0,0), (-1,-1), 4),
        ])
        tbl.setStyle(style)
        return tbl

    story = []
    blocks = re.split(r"\n\s*\n", text.strip())
    SHEET_RX = re.compile(r"^=+\s*Sheet:\s*(.+?)\s*=+$", re.I)

    for block in blocks:
        b = block.strip()
        if not b:
            continue
        if b == '---':
            story.append(PageBreak())
            continue

        lines = b.split('\n')

        # Heading
        m = re.match(r"^#{1,3}\s+(.+)$", b)
        if m:
            level = b.count("#", 0, b.find(" "))
            content = inline_markup(m.group(1))
            story.append(Paragraph(content, [h1, h2, h3][level-1]))
            story.append(Spacer(1, 8))
            continue

        # Sheet header
        ms = SHEET_RX.match(b)
        if ms:
            story.append(Paragraph(inline_markup(ms.group(0)), h2))
            story.append(Spacer(1, 6))
            continue

        # Lists
        if is_list_block(lines):
            bulletType = "bullet"
            numbered = all(re.match(r"\d+\.", ln.strip()) for ln in lines if ln.strip())
            items = [re.sub(r"^(?:\d+\.|[-*])\s*", "", ln.strip()) for ln in lines if ln.strip()]
            if numbered:
                bulletType = "1"
            story.append(ListFlowable([ListItem(Paragraph(inline_markup(it), base)) for it in items], bulletType=bulletType, leftIndent=1*cm))
            story.append(Spacer(1, 8))
            continue

        # Quotes
        if all(ln.strip().startswith(">") for ln in lines if ln.strip()):
            qtxt = "\n".join(ln.lstrip(">").strip() for ln in lines)
            story.append(Paragraph(inline_markup(qtxt), quote))
            story.append(Spacer(1, 6))
            continue

        # Table
        if is_table_block(lines):
            tbl = table_from_lines(lines)
            if tbl:
                story.append(tbl)
                story.append(Spacer(1, 8))
                continue

        # Paragraph
        story.append(Paragraph(inline_markup(b), base))
        story.append(Spacer(1, 6))

    doc.build(story)
